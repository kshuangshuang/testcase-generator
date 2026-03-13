import streamlit as st
from groq import Groq
import pandas as pd
import json
import io
from datetime import datetime
from pathlib import Path
import pypdf
from docx import Document

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="智能测试用例生成器",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 配置文件 ====================
CONFIG_FILE = "history.json"
MAX_HISTORY = 10  # 最多保存10条历史记录

# ==================== 辅助函数 ====================
def load_history():
    """加载历史记录"""
    if Path(CONFIG_FILE).exists():
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_history(history):
    """保存历史记录"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def add_to_history(doc_name, test_type, test_cases):
    """添加历史记录"""
    history = load_history()
    record = {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "doc_name": doc_name,
        "test_type": test_type,
        "test_cases": test_cases
    }
    history.insert(0, record)  # 最新的在前面
    if len(history) > MAX_HISTORY:
        history = history[:MAX_HISTORY]
    save_history(history)

def read_document(uploaded_file):
    """读取文档内容"""
    file_type = uploaded_file.name.split('.')[-1].lower()

    if file_type == 'pdf':
        # 读取 PDF
        pdf_reader = pypdf.PdfReader(uploaded_file)
        content = ""
        for page in pdf_reader.pages:
            content += page.extract_text() + "\n"
        return content

    elif file_type == 'docx':
        # 读取 Word 文档
        doc = Document(uploaded_file)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
            content += "\n"
        return content

    else:
        # 读取 TXT 或 MD
        return uploaded_file.getvalue().decode('utf-8')

def generate_test_cases(content, test_type, api_key):
    """调用 Groq API 生成测试用例"""
    client = Groq(api_key=api_key)

    # 根据测试类型构建提示词
    if test_type == "功能测试":
        prompt_instruction = """
你是一个专业的测试工程师。请基于以下需求文档，生成完整的功能测试用例。

要求：
1. **全面覆盖**：仔细分析需求文档，识别所有功能点、场景和业务流程，确保每个功能点都有对应的测试用例
2. **不限制数量**：根据需求的复杂程度生成测试用例，简单功能可能5-10条，复杂功能可能20-50条，不要人为限制在固定数量
3. **场景完整**：覆盖正常流程、基本异常流程、边界情况
4. 每个用例必须包含：用例ID、用例标题、前置条件、测试步骤、预期结果、优先级
5. 输出为 Markdown 表格格式
6. 所有内容使用中文

输出格式示例：
| 用例ID | 用例标题 | 前置条件 | 测试步骤 | 预期结果 | 优先级 |
|--------|---------|---------|---------|---------|--------|
| TC-001 | 用户登录 | 用户已注册 | 1. 打开登录页面<br>2. 输入用户名和密码<br>3. 点击登录按钮 | 登录成功，跳转到首页 | 高 |
"""
    elif test_type == "全面测试（含边界/异常）":
        prompt_instruction = """
你是一个专业的测试工程师。请基于以下需求文档，生成全面的测试用例。

要求：
1. **全面覆盖**：仔细分析需求文档，使用等价类划分、边界值分析、决策表等方法，识别所有测试点
2. **多维测试**：包含功能测试、边界值测试、异常场景测试、接口测试等
3. **不限制数量**：根据需求复杂程度生成，不要人为限制数量，确保测试覆盖完整
4. 每个用例必须包含：用例ID、用例标题、测试类型、前置条件、测试步骤、测试数据、预期结果、优先级
5. 输出为 Markdown 表格格式
6. 所有内容使用中文
7. **重要：只输出表格，不要有任何标题、分析说明或总结文字**

输出格式示例：
| 用例ID | 用例标题 | 测试类型 | 前置条件 | 测试步骤 | 测试数据 | 预期结果 | 优先级 |
|--------|---------|---------|---------|---------|---------|---------|--------|
"""
    else:  # 行业级专业测试
        prompt_instruction = """
你是一个资深的测试专家。请基于以下需求文档，生成符合 IEEE 829 标准的专业测试用例。

要求：
1. **全面覆盖**：从功能、边界、异常、性能、安全、兼容性等多维度分析需求，确保测试覆盖完整
2. **风险评估**：识别需求中的风险点和关键功能
3. **不限制数量**：根据需求复杂程度和风险评估，生成完整的测试用例集
4. 每个用例必须包含：用例ID、用例标题、测试类型、前置条件、测试步骤、测试数据、预期结果、优先级、状态、备注
5. 输出为 Markdown 表格格式
6. 在开头提供测试分析总结
7. 所有内容使用中文

输出格式示例：
## 测试分析总结
[简要分析需求，列出主要测试点和风险评估]

## 测试用例
| 用例ID | 用例标题 | 测试类型 | 前置条件 | 测试步骤 | 测试数据 | 预期结果 | 优先级 | 状态 | 备注 |
|--------|---------|---------|---------|---------|---------|---------|--------|------|------|
"""

    message_content = f"{prompt_instruction}\n\n## 需求文档内容\n\n{content}"

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": "你是一个专业的测试工程师，擅长编写高质量的测试用例。"
            },
            {
                "role": "user",
                "content": message_content
            }
        ],
        temperature=0.3,
        max_tokens=8000
    )

    return response.choices[0].message.content

def markdown_to_excel(markdown_content):
    """将 Markdown 表格转换为 Excel"""
    lines = markdown_content.split('\n')
    table_data = []
    headers = []

    for i, line in enumerate(lines):
        if line.startswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and any(cells):  # 确保不为空
                if not headers and all(cell and not cell.startswith('---') for cell in cells):
                    headers = cells
                elif not any(cell.startswith('---') for cell in cells):
                    table_data.append(cells)

    if table_data and headers:
        # 确保所有数据行的列数与表头一致
        max_cols = len(headers)
        table_data = [row + [''] * (max_cols - len(row)) if len(row) < max_cols else row[:max_cols] for row in table_data]

        df = pd.DataFrame(table_data, columns=headers)

        # 转换为 Excel
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='测试用例')
        output.seek(0)
        return output
    return None

# ==================== 主界面 ====================
st.title("📋 智能测试用例生成器")
st.markdown("---")

# 侧边栏
with st.sidebar:
    st.header("⚙️ 配置")

    # API Key 输入
    api_key = st.text_input(
        "Groq API Key",
        type="password",
        help="在 https://console.groq.com 获取 API Key",
        placeholder="gsk_xxxx..."
    )

    st.markdown("---")

    # 测试类型选择
    st.subheader("🎯 测试类型")
    test_type = st.selectbox(
        "选择测试类型",
        ["功能测试", "全面测试（含边界/异常）", "行业级专业测试"],
        label_visibility="collapsed"
    )

    # 测试类型说明
    if test_type == "功能测试":
        st.info("💡 覆盖正常流程和基本异常场景")
    elif test_type == "全面测试（含边界/异常）":
        st.info("💡 包含功能、边界值、异常等全方位测试")
    else:
        st.info("💡 符合 IEEE 829 标准的专业级测试用例")

    st.markdown("---")

    # 使用说明
    st.subheader("📖 使用说明")
    st.markdown("""
    ### 步骤
    1. 输入 Groq API Key
    2. 上传需求文档
    3. 选择测试类型
    4. 点击生成按钮
    5. 导出测试用例

    ### 支持格式
    - 📄 Word (.docx)
    - 📕 PDF (.pdf)
    - 📝 文本 (.txt)
    - 📃 Markdown (.md)

    ### 获取 API Key
    访问 [Groq Console](https://console.groq.com)
    注册账号后免费获取
    """)

# 主内容区域
col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📄 上传需求文档")

    uploaded_file = st.file_uploader(
        "上传需求文档",
        type=['txt', 'md', 'pdf', 'docx'],
        help="支持格式：Word、PDF、TXT、Markdown"
    )

    # 显示文档内容预览
    if uploaded_file:
        try:
            content = read_document(uploaded_file)
            st.success(f"✅ 已读取文件：{uploaded_file.name}")

            with st.expander("📖 查看文档内容", expanded=False):
                st.text_area("文档内容", content, height=300, disabled=True)
        except Exception as e:
            st.error(f"❌ 读取文件失败：{str(e)}")

with col2:
    st.subheader("🚀 生成测试用例")

    st.info(f"当前测试类型：**{test_type}**")

    generate_btn = st.button(
        "🚀 开始生成",
        type="primary",
        use_container_width=True,
        disabled=not uploaded_file or not api_key
    )

    if not uploaded_file:
        st.warning("⚠️ 请先上传需求文档")
    if not api_key:
        st.warning("⚠️ 请先输入 API Key")

# 生成测试用例
if generate_btn and uploaded_file and api_key:
    with st.spinner("🤖 AI 正在分析需求并生成测试用例，请稍候..."):
        try:
            # 读取文档
            content = read_document(uploaded_file)

            # 生成测试用例
            test_cases = generate_test_cases(content, test_type, api_key)

            # 保存到历史记录
            add_to_history(uploaded_file.name, test_type, test_cases)

            # 显示结果
            st.success("✅ 测试用例生成完成！")
            st.markdown("---")

            # 显示测试用例
            st.subheader("📊 测试用例预览")
            st.markdown(test_cases)

            # 导出功能
            st.markdown("---")
            st.subheader("💾 导出测试用例")

            col_export1, col_export2 = st.columns(2)

            # 导出 Markdown
            with col_export1:
                st.download_button(
                    label="📥 下载 Markdown 文档",
                    data=test_cases,
                    file_name=f"测试用例_{uploaded_file.name}.md",
                    mime="text/markdown",
                    use_container_width=True
                )

            # 导出 Excel
            with col_export2:
                try:
                    excel_data = markdown_to_excel(test_cases)
                    if excel_data:
                        st.download_button(
                            label="📥 下载 Excel 表格",
                            data=excel_data,
                            file_name=f"测试用例_{uploaded_file.name}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    else:
                        st.warning("⚠️ 未检测到表格格式，请使用 Markdown 导出")
                except Exception as e:
                    st.warning(f"⚠️ Excel 导出失败：{str(e)}")

        except Exception as e:
            st.error(f"❌ 生成失败：{str(e)}")
            st.info("💡 提示：请检查 API Key 是否正确")

# 历史记录
st.markdown("---")
st.subheader("📜 历史记录")

history = load_history()

if history:
    for i, record in enumerate(history):
        with st.expander(f"📄 {record['doc_name']} - {record['test_type']} ({record['time']})"):
            st.markdown(record['test_cases'])

            col_hist1, col_hist2 = st.columns(2)

            with col_hist1:
                st.download_button(
                    label="📥 下载 Markdown",
                    data=record['test_cases'],
                    file_name=f"测试用例_{record['doc_name']}.md",
                    mime="text/markdown",
                    key=f"hist_md_{i}"
                )

            with col_hist2:
                try:
                    excel_data = markdown_to_excel(record['test_cases'])
                    if excel_data:
                        st.download_button(
                            label="📥 下载 Excel",
                            data=excel_data,
                            file_name=f"测试用例_{record['doc_name']}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"hist_xlsx_{i}"
                        )
                except:
                    pass
else:
    st.info("📭 暂无历史记录，生成测试用例后会自动保存")

# 页脚
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px;'>
    <p>由 Groq + Llama 3.1 驱动 | 永久免费使用</p>
    <p>需要 API Key？访问 <a href='https://console.groq.com' target='_blank'>Groq Console</a></p>
</div>
""", unsafe_allow_html=True)
